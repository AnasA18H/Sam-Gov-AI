"""
Document Text and Data Extractor
Implements ensemble method with multiple extractors for maximum accuracy
"""
import re
import logging
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Any
import json

try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

try:
    import camelot
    CAMELOT_AVAILABLE = True
except ImportError:
    CAMELOT_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from openpyxl import load_workbook
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False

from ..core.config import settings

logger = logging.getLogger(__name__)


class DocumentExtractor:
    """
    Robust document extractor using ensemble method for maximum accuracy.
    Tries multiple extractors and returns the best quality result.
    """
    
    def __init__(self):
        self.debug_extracts_dir = settings.DEBUG_EXTRACTS_DIR
        
    def extract_text_robustly(self, file_path: str, opportunity_id: int = None, document_id: int = None) -> Dict[str, Any]:
        """
        Extract text from document using ensemble method.
        Tries multiple extractors and returns the best quality result.
        
        Args:
            file_path: Path to the document file
            opportunity_id: Optional opportunity ID for debug file organization
            document_id: Optional document ID for debug file naming
            
        Returns:
            Dictionary with:
            - 'text': Extracted text
            - 'quality_score': Quality score (0-100)
            - 'method_used': Name of the extractor that provided the best result
            - 'all_results': List of all extraction attempts with scores
        """
        file_path_obj = Path(file_path)
        
        # Handle relative paths
        if not file_path_obj.is_absolute():
            project_root = settings.PROJECT_ROOT
            abs_path = project_root / file_path
            if abs_path.exists():
                file_path_obj = abs_path
            elif hasattr(settings, 'STORAGE_BASE_PATH'):
                storage_base = Path(settings.STORAGE_BASE_PATH)
                if storage_base.is_absolute():
                    abs_path = storage_base.parent / file_path if 'backend/data' in str(file_path) else storage_base / file_path
                else:
                    abs_path = project_root / storage_base.parent / file_path if 'backend/data' in str(file_path) else project_root / storage_base / file_path
                if abs_path.exists():
                    file_path_obj = abs_path
        
        if not file_path_obj.exists():
            raise FileNotFoundError(f"Document file not found: {file_path}")
        
        file_ext = file_path_obj.suffix.lower()
        
        # Route to appropriate extractor based on file type
        if file_ext == '.pdf':
            return self._extract_pdf_robustly(file_path_obj, opportunity_id, document_id)
        elif file_ext in ['.doc', '.docx']:
            return self._extract_word_robustly(file_path_obj, opportunity_id, document_id)
        elif file_ext in ['.xls', '.xlsx']:
            return self._extract_excel_robustly(file_path_obj, opportunity_id, document_id)
        elif file_ext == '.txt':
            return self._extract_txt(file_path_obj, opportunity_id, document_id)
        else:
            logger.warning(f"Unsupported file type: {file_ext}")
            return {
                'text': '',
                'quality_score': 0,
                'method_used': 'unsupported',
                'all_results': []
            }
    
    def _extract_pdf_robustly(self, pdf_path: Path, opportunity_id: int = None, document_id: int = None) -> Dict[str, Any]:
        """Extract text from PDF using ensemble method"""
        results = []
        
        # Try PyMuPDF first (fast, good for text)
        if PYMUPDF_AVAILABLE:
            try:
                text = self._extract_with_pymupdf(pdf_path)
                score = self._quality_score(text)
                results.append({
                    'method': 'pymupdf',
                    'text': text,
                    'quality_score': score,
                    'length': len(text)
                })
                logger.info(f"PyMuPDF extraction: {len(text)} chars, quality: {score:.2f}")
            except Exception as e:
                logger.warning(f"PyMuPDF extraction failed: {str(e)}")
        
        # Try PDFPlumber (best for tables and precision)
        if PDFPLUMBER_AVAILABLE:
            try:
                text = self._extract_with_pdfplumber(pdf_path)
                score = self._quality_score(text)
                results.append({
                    'method': 'pdfplumber',
                    'text': text,
                    'quality_score': score,
                    'length': len(text)
                })
                logger.info(f"PDFPlumber extraction: {len(text)} chars, quality: {score:.2f}")
            except Exception as e:
                logger.warning(f"PDFPlumber extraction failed: {str(e)}")
        
        # Select best result
        if results:
            best_result = max(results, key=lambda x: x['quality_score'])
            
            # Save debug extract
            if opportunity_id:
                self._save_debug_extract(
                    opportunity_id=opportunity_id,
                    document_id=document_id,
                    filename=pdf_path.name,
                    text=best_result['text'],
                    method=best_result['method'],
                    quality_score=best_result['quality_score'],
                    all_results=results
                )
            
            return {
                'text': best_result['text'],
                'quality_score': best_result['quality_score'],
                'method_used': best_result['method'],
                'all_results': results
            }
        
        # Fallback: empty result
        logger.error(f"All PDF extraction methods failed for {pdf_path}")
        return {
            'text': '',
            'quality_score': 0,
            'method_used': 'none',
            'all_results': []
        }
    
    def _extract_with_pymupdf(self, pdf_path: Path) -> str:
        """Extract text using PyMuPDF with structure preservation"""
        doc = fitz.open(pdf_path)
        text_blocks = []
        
        for page_num, page in enumerate(doc):
            # Get structured text blocks
            blocks = page.get_text("dict")["blocks"]
            for block in blocks:
                if "lines" in block:  # Text block
                    block_text = ""
                    for line in block["lines"]:
                        for span in line["spans"]:
                            # Preserve font info for headers (bold text)
                            if span.get("flags", 0) & 2**4:  # Bold
                                block_text += f"**{span['text']}** "
                            else:
                                block_text += span["text"] + " "
                    if block_text.strip():
                        text_blocks.append(block_text.strip())
        
        doc.close()
        return "\n".join(text_blocks)
    
    def _extract_with_pdfplumber(self, pdf_path: Path) -> str:
        """Extract text using PDFPlumber with table detection"""
        text_parts = []
        
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                # Extract regular text
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
                
                # Extract tables as formatted text
                tables = page.extract_tables()
                if tables:
                    for table in tables:
                        if table and len(table) > 1:
                            table_text = self._format_table_as_text(table)
                            text_parts.append(f"\n[TABLE]\n{table_text}\n[/TABLE]\n")
        
        return "\n\n".join(text_parts)
    
    def _extract_word_robustly(self, doc_path: Path, opportunity_id: int = None, document_id: int = None) -> Dict[str, Any]:
        """Extract text from Word document"""
        if not DOCX_AVAILABLE:
            return {
                'text': '',
                'quality_score': 0,
                'method_used': 'none',
                'all_results': []
            }
        
        try:
            doc = DocxDocument(doc_path)
            text_parts = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract tables
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    if row_text.strip():
                        table_text.append(row_text)
                if table_text:
                    text_parts.append(f"\n[TABLE]\n" + "\n".join(table_text) + "\n[/TABLE]\n")
            
            text = "\n".join(text_parts)
            score = self._quality_score(text)
            
            # Save debug extract
            if opportunity_id:
                self._save_debug_extract(
                    opportunity_id=opportunity_id,
                    document_id=document_id,
                    filename=doc_path.name,
                    text=text,
                    method='python-docx',
                    quality_score=score,
                    all_results=[{'method': 'python-docx', 'text': text, 'quality_score': score, 'length': len(text)}]
                )
            
            return {
                'text': text,
                'quality_score': score,
                'method_used': 'python-docx',
                'all_results': [{'method': 'python-docx', 'text': text, 'quality_score': score, 'length': len(text)}]
            }
        except Exception as e:
            logger.error(f"Error extracting Word document {doc_path}: {str(e)}")
            return {
                'text': '',
                'quality_score': 0,
                'method_used': 'none',
                'all_results': []
            }
    
    def _extract_excel_robustly(self, excel_path: Path, opportunity_id: int = None, document_id: int = None) -> Dict[str, Any]:
        """Extract text from Excel file"""
        if not OPENPYXL_AVAILABLE:
            return {
                'text': '',
                'quality_score': 0,
                'method_used': 'none',
                'all_results': []
            }
        
        try:
            workbook = load_workbook(excel_path, data_only=True)
            text_parts = []
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                sheet_text = []
                
                for row in sheet.iter_rows(values_only=True):
                    row_text = " | ".join([str(cell) if cell is not None else "" for cell in row])
                    if row_text.strip():
                        sheet_text.append(row_text)
                
                if sheet_text:
                    text_parts.append(f"\n--- Sheet: {sheet_name} ---\n" + "\n".join(sheet_text))
            
            text = "\n\n".join(text_parts)
            score = self._quality_score(text)
            
            # Save debug extract
            if opportunity_id:
                self._save_debug_extract(
                    opportunity_id=opportunity_id,
                    document_id=document_id,
                    filename=excel_path.name,
                    text=text,
                    method='openpyxl',
                    quality_score=score,
                    all_results=[{'method': 'openpyxl', 'text': text, 'quality_score': score, 'length': len(text)}]
                )
            
            return {
                'text': text,
                'quality_score': score,
                'method_used': 'openpyxl',
                'all_results': [{'method': 'openpyxl', 'text': text, 'quality_score': score, 'length': len(text)}]
            }
        except Exception as e:
            logger.error(f"Error extracting Excel file {excel_path}: {str(e)}")
            return {
                'text': '',
                'quality_score': 0,
                'method_used': 'none',
                'all_results': []
            }
    
    def _extract_txt(self, txt_path: Path, opportunity_id: int = None, document_id: int = None) -> Dict[str, Any]:
        """Extract text from plain text file"""
        try:
            with open(txt_path, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
            
            score = self._quality_score(text)
            
            # Save debug extract
            if opportunity_id:
                self._save_debug_extract(
                    opportunity_id=opportunity_id,
                    document_id=document_id,
                    filename=txt_path.name,
                    text=text,
                    method='plain_text',
                    quality_score=score,
                    all_results=[{'method': 'plain_text', 'text': text, 'quality_score': score, 'length': len(text)}]
                )
            
            return {
                'text': text,
                'quality_score': score,
                'method_used': 'plain_text',
                'all_results': [{'method': 'plain_text', 'text': text, 'quality_score': score, 'length': len(text)}]
            }
        except Exception as e:
            logger.error(f"Error extracting text file {txt_path}: {str(e)}")
            return {
                'text': '',
                'quality_score': 0,
                'method_used': 'none',
                'all_results': []
            }
    
    def extract_clin_tables(self, pdf_path: Path) -> List[Dict[str, Any]]:
        """
        Specialized CLIN extraction from SF1449 forms and similar documents.
        Uses camelot for clean tables, falls back to pdfplumber.
        """
        clin_tables = []
        
        # METHOD 1: Camelot for clean tables
        if CAMELOT_AVAILABLE:
            try:
                tables = camelot.read_pdf(
                    str(pdf_path),
                    pages='all',
                    flavor='lattice',  # For grid-like tables
                    strip_text='\n',
                    edge_tol=500
                )
                
                for table in tables:
                    df = table.df
                    if df.empty:
                        continue
                    
                    # Check if this looks like a CLIN table
                    first_row_text = df.iloc[0].to_string().upper() if len(df) > 0 else ""
                    if any(header in first_row_text for header in ['CLIN', 'ITEM NO', 'LINE ITEM', 'ITEM NUMBER']):
                        # Convert to list of dicts
                        headers = [str(h).strip() for h in df.iloc[0].tolist()] if len(df) > 0 else []
                        for idx in range(1, len(df)):
                            row = df.iloc[idx]
                            row_dict = {headers[i]: str(row.iloc[i]).strip() if i < len(headers) else "" 
                                       for i in range(len(row))}
                            clin_tables.append(row_dict)
                        
                        logger.info(f"Found CLIN table with {len(df)-1} rows using Camelot")
                        break
            except Exception as e:
                logger.warning(f"Camelot CLIN extraction failed: {str(e)}")
        
        # METHOD 2: PDFPlumber as fallback
        if not clin_tables and PDFPLUMBER_AVAILABLE:
            try:
                with pdfplumber.open(pdf_path) as pdf:
                    for page in pdf.pages:
                        tables = page.extract_tables()
                        for table in tables:
                            if table and len(table) > 1:
                                first_row = ' '.join(str(cell) for cell in table[0] if cell).upper()
                                if any(keyword in first_row for keyword in ['CLIN', 'ITEM', 'LINE']):
                                    headers = [str(cell).strip() if cell else "" for cell in table[0]]
                                    for row in table[1:]:
                                        if row:
                                            row_dict = {headers[i]: str(row[i]).strip() if i < len(headers) and row[i] else "" 
                                                       for i in range(min(len(row), len(headers)))}
                                            clin_tables.append(row_dict)
                                    logger.info(f"Found CLIN table with {len(table)-1} rows using PDFPlumber")
                                    break
            except Exception as e:
                logger.warning(f"PDFPlumber CLIN extraction failed: {str(e)}")
        
        return clin_tables
    
    def extract_sf1449_fields(self, text: str) -> Dict[str, str]:
        """Targeted extraction for Standard Form 1449"""
        patterns = {
            'solicitation_number': r'SOLICITATION\s+NUMBER\s*[:]?\s*([A-Z0-9-]+)',
            'response_deadline': r'OFFER\s+DUE\s+DATE[/]?\s*([A-Za-z]+\s+\d{1,2},?\s+\d{4}\s+[\d:]+\s*[APM]+\s*[A-Z]+)',
            'delivery_address': r'DELIVER\s+TO.*?\n(.*?)(?=\n\s*\n|$)',
        }
        
        extracted = {}
        for field, pattern in patterns.items():
            match = re.search(pattern, text, re.IGNORECASE | re.DOTALL | re.MULTILINE)
            if match:
                extracted[field] = self._clean_text(match.group(1))
        
        return extracted
    
    def _quality_score(self, text: str) -> float:
        """
        Score extraction quality (0-100).
        Higher score = better quality.
        """
        if not text or len(text.strip()) == 0:
            return 0.0
        
        score = 50.0  # Base score
        
        # Length factor (longer is generally better, but not too long)
        length = len(text)
        if 100 <= length <= 100000:  # Sweet spot
            score += 20
        elif length > 100000:
            score += 10
        elif length < 100:
            score -= 20
        
        # Alphanumeric ratio (higher is better)
        alphanum_chars = sum(c.isalnum() for c in text)
        total_chars = len(text)
        if total_chars > 0:
            alphanum_ratio = alphanum_chars / total_chars
            score += alphanum_ratio * 20
        
        # Check for common PDF artifacts (negative)
        if re.search(r'\(cid:\d+\)', text):
            score -= 10  # PDF encoding artifacts
        
        if re.search(r'[^\x20-\x7E\n\r\t]', text[:1000]):
            score -= 5  # Non-printable characters
        
        # Check for structured content (positive)
        if re.search(r'\n{2,}', text):
            score += 5  # Has paragraphs
        
        if re.search(r'[A-Z]{2,}', text[:500]):
            score += 5  # Has uppercase words (headers, etc.)
        
        # Ensure score is within bounds
        return max(0.0, min(100.0, score))
    
    def _clean_text(self, text: str) -> str:
        """Clean extracted text"""
        if not text:
            return ""
        
        # Remove PDF encoding artifacts
        text = re.sub(r'\(cid:\d+\)', '', text)
        
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove non-printable characters except newlines/tabs
        text = re.sub(r'[^\x20-\x7E\n\r\t]', '', text)
        
        return text.strip()
    
    def _format_table_as_text(self, table: List[List]) -> str:
        """Format extracted table as readable text"""
        if not table or len(table) == 0:
            return ""
        
        rows = []
        for row in table:
            # Filter out None values and convert to strings
            row_text = " | ".join([str(cell).strip() if cell else "" for cell in row])
            if row_text.strip():
                rows.append(row_text)
        
        return "\n".join(rows)
    
    def _save_debug_extract(
        self, 
        opportunity_id: int, 
        document_id: int = None,
        filename: str = None,
        text: str = None,
        method: str = None,
        quality_score: float = None,
        all_results: List[Dict] = None
    ):
        """
        Save debug extraction results to backend/data/debug_extracts/
        """
        try:
            debug_dir = self.debug_extracts_dir / f"opportunity_{opportunity_id}"
            debug_dir.mkdir(parents=True, exist_ok=True)
            
            # Create debug filename
            if document_id and filename:
                debug_filename = f"{document_id}_{filename}_extracted.txt"
            elif filename:
                debug_filename = f"{filename}_extracted.txt"
            else:
                debug_filename = f"document_{document_id}_extracted.txt"
            
            debug_file = debug_dir / debug_filename
            
            # Write extraction results
            with open(debug_file, 'w', encoding='utf-8') as f:
                f.write("=" * 80 + "\n")
                f.write("DOCUMENT TEXT EXTRACTION DEBUG\n")
                f.write("=" * 80 + "\n\n")
                
                f.write(f"Document: {filename or 'Unknown'}\n")
                f.write(f"Document ID: {document_id or 'N/A'}\n")
                f.write(f"Opportunity ID: {opportunity_id}\n")
                f.write(f"Extraction Method: {method or 'N/A'}\n")
                f.write(f"Quality Score: {quality_score:.2f}\n" if quality_score else "Quality Score: N/A\n")
                f.write(f"Text Length: {len(text)} characters\n" if text else "Text Length: 0\n")
                
                if all_results:
                    f.write("\n" + "=" * 80 + "\n")
                    f.write("ALL EXTRACTION ATTEMPTS:\n")
                    f.write("=" * 80 + "\n")
                    for result in all_results:
                        f.write(f"\nMethod: {result.get('method', 'unknown')}\n")
                        f.write(f"  Quality Score: {result.get('quality_score', 0):.2f}\n")
                        f.write(f"  Length: {result.get('length', 0)} characters\n")
                
                if text:
                    f.write("\n" + "=" * 80 + "\n")
                    f.write("EXTRACTED TEXT:\n")
                    f.write("=" * 80 + "\n\n")
                    f.write(text)
                
                # Also save JSON metadata
                metadata_file = debug_dir / f"{document_id}_{filename}_metadata.json" if document_id and filename else debug_dir / f"metadata_{document_id}.json"
                metadata = {
                    'document_id': document_id,
                    'filename': filename,
                    'opportunity_id': opportunity_id,
                    'method_used': method,
                    'quality_score': quality_score,
                    'text_length': len(text) if text else 0,
                    'all_results': [
                        {
                            'method': r.get('method'),
                            'quality_score': r.get('quality_score'),
                            'length': r.get('length')
                        } for r in (all_results or [])
                    ]
                }
                
                with open(metadata_file, 'w', encoding='utf-8') as mf:
                    json.dump(metadata, mf, indent=2)
            
            logger.info(f"Saved debug extract to {debug_file}")
        except Exception as e:
            logger.warning(f"Failed to save debug extract: {str(e)}")
